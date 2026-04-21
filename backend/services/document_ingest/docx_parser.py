"""DOCX parser — extracts paragraphs (with heading levels) and tables.

Requires the ``python-docx`` package. If the dependency is missing at
runtime, parsing raises a clear ImportError so the operator can install it
without having to read a stack trace.
"""

from __future__ import annotations

import io
from typing import List

from api.schemas.document import Section, Table

from .base import ParserResult, register


def _heading_level(style_name: str) -> int:
    """Return 1-6 for Word's ``Heading 1`` … ``Heading 6`` styles, else 0."""
    if not style_name:
        return 0
    name = style_name.strip().lower()
    if name.startswith("heading "):
        try:
            level = int(name.split(" ", 1)[1])
        except ValueError:
            return 0
        return level if 1 <= level <= 6 else 0
    return 0


@register("docx")
def parse_docx(data: bytes) -> ParserResult:
    try:
        import docx  # type: ignore
    except ImportError as exc:  # pragma: no cover — exercised only in misconfigured envs
        raise ImportError(
            "python-docx is required to parse .docx files. "
            "Install it with: pip install python-docx"
        ) from exc

    document = docx.Document(io.BytesIO(data))

    sections: List[Section] = []
    text_lines: list[str] = []

    current = Section(heading="", body="", level=0)
    body_buf: list[str] = []

    def _flush() -> None:
        if current.heading or body_buf:
            current.body = "\n".join(body_buf).strip()
            sections.append(current.model_copy())

    for para in document.paragraphs:
        line = para.text or ""
        level = _heading_level(getattr(para.style, "name", "") or "")
        if level > 0:
            _flush()
            current.heading = line.strip()
            current.level = level
            body_buf = []
            text_lines.append(line)
            continue
        if line:
            body_buf.append(line)
            text_lines.append(line)
    _flush()

    tables: List[Table] = []
    for idx, tbl in enumerate(document.tables):
        rows = [[(cell.text or "").strip() for cell in row.cells] for row in tbl.rows]
        if not rows:
            continue
        headers, body = rows[0], rows[1:]
        tables.append(Table(name=f"table_{idx + 1}", headers=headers, rows=body))

    return "\n".join(text_lines).strip(), sections, tables
